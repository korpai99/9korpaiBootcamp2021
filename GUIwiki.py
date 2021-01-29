# GUIwiki.py
import wikipedia

# python to docx
from docx import Document
def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)
    # Summary สำหรับบทความที่สรุป
    data = wikipedia.summary(keyword)

    # Page + content บทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content

    doc = Document() #สร้างไฟล์ word ใน python
    doc.add_heading(keyword,0)

    doc.add_paragraph(data2)
    doc.save(keyword+ '.docx')
    print('สร้างไฟล์สำเร็จ')

# เปลี่ยนเป็นภาษาไทย
wikipedia.set_lang('th')
from tkinter import *
from tkinter import ttk
from tkinter import messagebox

GUI = Tk()
GUI.title('โปรแกรม Wiki By 9korpai')
GUI.geometry('400x200')#ขนาดกล่อง popup

# Config
FONT1 = ('Angsana New',15)

# คำอธิบาย
L = ttk.Label(GUI, text='ค้นหาบทความ',font=FONT1)
L.pack(pady=10)

# ช่องค้นหาข้อมูล
v_search = StringVar()#กล่องสำหรับเก็บคีย์เวิร์ด
E1 = ttk.Entry(GUI,textvariable = v_search,font=FONT1,width=30)
E1.pack(pady=10)

# ปุ่มค้นหา
def Search():
    keyword = v_search.get() # .get() คือ ดึงข้อมูลเข้ามา
    try:
        #ลองค้าหาดูว่าได้ผลลัพท์หรือไม่้ หากได้ให้ผ่านไป
        language= v_radio.get() # th / en / zh
        Wiki(keyword,language)
        messagebox.showinfo('บันทึกสำเร็จ','ค้นหาข้อความสำเร็จ บันทึกเรียบร้อยแล้ว')
    except:
        # หากรันคำสั่งแล้วมีปัญหา แสดงข้อความแจ้งเตือน
        messagebox.showwarning('Keywoed Error','กรุณากรอกคำค้นหาใหม่')

        
    # print(wikipedia.search(keyword))# บอกจำนวนคีย์เวิดร์ว่ามีกี่คำ
    # result = wikipedia.summary(keyword)
    # print(result)


# ปุ่มค้นหา
B1 = ttk.Button(GUI,text='Search',command=Search)# command คือ ลิงค์ปุ่มการค้นหา
B1.pack(ipadx=20,ipady=10) #ipadx ขยายภายในปุ่มแนวนอน

# เมนูเลือกภาษา
F1 = Frame(GUI)
F1.pack()

v_radio = StringVar()# ช่องเก็บข้อมูลภาษา

RB1 = ttk.Radiobutton(F1,text='ภาษาไทย',variable=v_radio,value='th')
RB2 = ttk.Radiobutton(F1,text='ภาษาอังกฤษ',variable=v_radio,value='en')
RB3 = ttk.Radiobutton(F1,text='ภาษาจีน',variable=v_radio,value='zh')
RB1.invoke() #สั่งให้ค่าเริ่มต้นเป็นภาษาไทย

RB1.grid(row=0,column=0)
RB2.grid(row=0,column=1)
RB3.grid(row=0,column=2)

GUI.mainloop()
